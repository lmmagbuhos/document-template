#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE_LOGO = ROOT / "resources/brand/logos/viber_image_2026-05-19_17-29-23-542.jpg"
TRIMMED_LOGO = ROOT / "resources/brand/logos/oak-drive-logo-trimmed.png"
OUTPUT = ROOT / "resources/templates/docx/oak-drive-policy-template.docx"


def crop_logo() -> None:
    image = Image.open(SOURCE_LOGO).convert("RGB")
    bg = Image.new("RGB", image.size, image.getpixel((0, 0)))
    diff = Image.eval(Image.blend(image, bg, 0.0), lambda px: px)

    # Use a threshold against the white background to trim scanner/export whitespace.
    pixels = image.load()
    width, height = image.size
    xs = []
    ys = []
    for y in range(height):
        for x in range(width):
            r, g, b = pixels[x, y]
            if min(abs(255 - r), abs(255 - g), abs(255 - b)) > 8 and (r < 248 or g < 248 or b < 248):
                xs.append(x)
                ys.append(y)

    if xs and ys:
        box = (max(min(xs) - 8, 0), max(min(ys) - 8, 0), min(max(xs) + 8, width), min(max(ys) + 8, height))
        image = image.crop(box)

    image.save(TRIMMED_LOGO)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_paragraph_bottom_border(paragraph, size="24", color="000000"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_style_font(style, name="Arial", size=11, bold=False):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_run(paragraph, text, bold=False, size=11):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(1.72)
    section.bottom_margin = Inches(0.55)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        set_style_font(styles[style_name], size=11, bold=True)
        styles[style_name].paragraph_format.space_before = Pt(6)
        styles[style_name].paragraph_format.space_after = Pt(8)
        styles[style_name].paragraph_format.keep_with_next = True

    return doc


def build_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    table = header.add_table(rows=1, cols=2, width=Inches(7.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.65)
    table.columns[1].width = Inches(1.75)

    for cell in table.row_cells(0):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for side in ("top", "left", "bottom", "right"):
            set_cell_border(cell, **{side: {"val": "nil"}})

    left = table.cell(0, 0)
    left.text = ""
    lines = [
        ("Policy Title: ", "{{ policy_title }}", True),
        ("Policy Number: ", "{{ policy_number }}", False),
        ("Issuing Department: ", "{{ issuing_department }}", False),
        ("Supersedes: ", "{{ supersedes }}", False),
        ("Effective Date: ", "{{ effective_date }}", False),
    ]
    for index, (label, value, value_bold) in enumerate(lines):
        paragraph = left.paragraphs[0] if index == 0 else left.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(paragraph, label, size=11)
        add_run(paragraph, value, bold=value_bold, size=11)

    right = table.cell(0, 1)
    right.text = ""
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(TRIMMED_LOGO), width=Inches(1.22))

    divider = header.add_paragraph()
    divider.paragraph_format.space_before = Pt(8)
    divider.paragraph_format.space_after = Pt(0)
    set_paragraph_bottom_border(divider, size="24")


def add_sample_body(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    add_run(p, "1.0 Policy Objective", bold=True)

    p = doc.add_paragraph()
    add_run(
        p,
        "{{ policy_objective }}",
    )

    p = doc.add_paragraph(style="Heading 1")
    add_run(p, "2.0 Scope and Coverage", bold=True)

    p = doc.add_paragraph()
    add_run(p, "{{ scope_and_coverage }}")

    p = doc.add_paragraph(style="Heading 1")
    add_run(p, "3.0 Procedures and Guidelines", bold=True)

    p = doc.add_paragraph()
    add_run(p, "{{ procedures_and_guidelines }}")

    p = doc.add_paragraph(style="Heading 2")
    add_run(p, "Approval Process", bold=True)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.15)
    table.columns[1].width = Inches(2.05)
    headers = ["Employee Level", "Approving Authority"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = ""
        para = cell.paragraphs[0]
        add_run(para, label, bold=True)
    set_repeat_table_header(table.rows[0])
    sample_rows = [
        ("{{ employee_level_1 }}", "{{ approving_authority_1 }}"),
        ("{{ employee_level_2 }}", "{{ approving_authority_2 }}"),
    ]
    for row_idx, row_values in enumerate(sample_rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            add_run(cell.paragraphs[0], value)
    for row in table.rows:
        for cell in row.cells:
            for side in ("top", "left", "bottom", "right"):
                set_cell_border(cell, **{side: {"val": "single", "sz": "8", "color": "000000", "space": "0"}})

    doc.add_paragraph()
    p = doc.add_paragraph(style="Heading 1")
    add_run(p, "4.0 Compliance and Accountability", bold=True)
    p = doc.add_paragraph()
    add_run(p, "{{ compliance_and_accountability }}")

    p = doc.add_paragraph(style="Heading 1")
    add_run(p, "5.0 Related Policies", bold=True)
    p = doc.add_paragraph()
    add_run(p, "{{ related_policies }}")

    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.45)
    sig_table.columns[1].width = Inches(3.45)
    labels = [
        ("Prepared by:", "{{ prepared_by_name }}", "{{ prepared_by_title }}"),
        ("Approved by:", "{{ approved_by_name }}", "{{ approved_by_title }}"),
    ]
    for idx, (label, name, title) in enumerate(labels):
        cell = sig_table.cell(0, idx)
        cell.text = ""
        for side in ("top", "left", "bottom", "right"):
            set_cell_border(cell, **{side: {"val": "nil"}})
        add_run(cell.paragraphs[0], label)
        spacer = cell.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)
        add_run(spacer, "{{ signature_image }}")
        name_p = cell.add_paragraph()
        name_p.paragraph_format.space_after = Pt(0)
        add_run(name_p, name, bold=True)
        title_p = cell.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        add_run(title_p, title)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    crop_logo()
    doc = configure_document()
    build_header(doc)
    add_sample_body(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
